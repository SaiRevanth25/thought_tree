"""File processing service for extracting text and chunking documents."""

import io
from typing import BinaryIO

import structlog
import tiktoken

from core.config import settings

logger = structlog.getLogger(__name__)


class FileProcessor:
    """Service for processing uploaded files and extracting text content."""

    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "text/plain": "txt",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/markdown": "md",
    }

    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        # Use cl100k_base encoding (used by OpenAI embeddings)
        self.tokenizer = tiktoken.get_encoding("cl100k_base")

    async def extract_text(self, file_content: bytes, mime_type: str) -> str:
        """
        Extract text from a file based on its MIME type.

        Args:
            file_content: Raw file bytes
            mime_type: MIME type of the file

        Returns:
            Extracted text content
        """
        file_type = self.SUPPORTED_TYPES.get(mime_type)

        if file_type == "pdf":
            return await self._extract_pdf(file_content)
        elif file_type == "txt" or file_type == "md":
            return await self._extract_text_file(file_content)
        elif file_type == "docx":
            return await self._extract_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")

    async def _extract_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_content))
            text_parts = []

            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num}]\n{page_text}")

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to extract PDF: {e}")
            raise ValueError(f"Failed to extract PDF content: {str(e)}")

    async def _extract_text_file(self, file_content: bytes) -> str:
        """Extract text from plain text or markdown file."""
        try:
            # Try UTF-8 first, then fall back to other encodings
            for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode file with any supported encoding")
        except Exception as e:
            logger.error(f"Failed to extract text file: {e}")
            raise ValueError(f"Failed to extract text content: {str(e)}")

    async def _extract_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_content))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to extract DOCX: {e}")
            raise ValueError(f"Failed to extract DOCX content: {str(e)}")

    def chunk_text(self, text: str) -> list[str]:
        """
        Split text into overlapping chunks based on token count.

        Args:
            text: Full text content to chunk

        Returns:
            List of text chunks
        """
        if not text.strip():
            return []

        # Tokenize the entire text
        tokens = self.tokenizer.encode(text)
        total_tokens = len(tokens)

        if total_tokens <= self.chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < total_tokens:
            # Get chunk tokens
            end = min(start + self.chunk_size, total_tokens)
            chunk_tokens = tokens[start:end]

            # Decode tokens back to text
            chunk_text = self.tokenizer.decode(chunk_tokens)
            chunks.append(chunk_text.strip())

            # Move start position with overlap
            start = end - self.chunk_overlap

            # Prevent infinite loop at the end
            if start >= total_tokens - self.chunk_overlap:
                break

        logger.info(
            f"Chunked text into {len(chunks)} chunks "
            f"(total tokens: {total_tokens}, chunk_size: {self.chunk_size})"
        )

        return chunks

    def count_tokens(self, text: str) -> int:
        """Count the number of tokens in a text string."""
        return len(self.tokenizer.encode(text))

    def is_supported_type(self, mime_type: str) -> bool:
        """Check if the MIME type is supported."""
        return mime_type in self.SUPPORTED_TYPES


# Singleton instance
file_processor = FileProcessor()
